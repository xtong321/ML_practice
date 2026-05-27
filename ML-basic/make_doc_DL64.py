"""
generate docs file from pdf file
"""

import os

try:
    import docx
except ImportError:
    # 自动安装所需的 python-docx 库
    os.system('pip install python-docx')
    import docx

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# 设置页面边距
sections = doc.sections
for section in sections:
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)

# 样式设置助手
def add_p(text, bold=False, size=11, align=WD_ALIGN_PARAGRAPH.LEFT, space_after=6, italic=False):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    run.font.name = 'Arial'
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    return p

# --- PAGE 1 ---
add_p("Application for Change of Address or Replacement on valid Texas Driver License (DL), Commercial Driver License (CDL) & Identification Card (ID)", bold=True, size=14, align=WD_ALIGN_PARAGRAPH.CENTER)
add_p("Form: DL-64 01/2025", size=10, align=WD_ALIGN_PARAGRAPH.CENTER, italic=True)

p_box = doc.add_paragraph()
p_box.paragraph_format.space_after = Pt(12)
r_box = p_box.add_run("⚠️ DO NOT MAIL CASH (请勿邮寄现金)\nMAIL COMPLETED FORM AND $10 FEE TO:\nTexas Department of Public Safety, PO Box 149008, Austin, Texas, 78714-9008\n(Make check or money order payable to: Texas Department of Public Safety)")
r_box.font.name = 'Arial'
r_box.font.size = Pt(11)
r_box.bold = True

add_p("Are you a citizen of the United States?   [  ] Yes   [  ] No", bold=True)
add_p("Would you like to register as an organ donor?   [  ] Yes   [  ] No", bold=True)

fields = [
    "Driver License Number (驾照号): ________________________",
    "Identification Card Number (ID卡号): ____________________",
    "Date of Birth (出生日期 MM/DD/YYYY): _____ / _____ / ________",
    "Phone Number (联系电话 - 可填您的美国手机): ________________________",
    "Social Security Number (社会安全号 SSN): _____ - _____ - _________"
]
for f in fields:
    add_p(f, space_after=8)

add_p("Applicant Information (申请人信息)", bold=True, size=12, space_after=8)
add_p("Last Name (姓): _______________________________________")
add_p("First Name (名): ______________________________________")
add_p("Middle Name / Birth Surname (中间名): ___________________")
add_p("Suffix (后缀如 SR., JR. 等): ___________")
add_p("Email (电子邮箱): ______________________________________")

add_p("Residence Address (居住地址 - 必须是实际住址，不能填PO Box)", bold=True, size=12, space_after=8)
add_p("Street Address (新公寓街区与门牌号): __________________________________________________")
add_p("City (城市): ___________________    County (郡/县): _________________")
add_p("State (州): TX    Zip Code (邮编): _________________")

add_p("Mailing Address (邮寄地址 - 新驾照将寄到此地址)", bold=True, size=12, space_after=8)
add_p("Street Address or P.O. Box (若与上方相同可照抄): _________________________________________________")
add_p("City (城市): ___________________    County (郡/县): _________________")
add_p("State (州): TX    Zip Code (邮编): _________________")

add_p("---------------------------------- PAGE 2 ----------------------------------", align=WD_ALIGN_PARAGRAPH.CENTER, space_after=18)

# --- PAGE 2 ---
add_p("Voter Registration Application (选民登记选项 - 非美国公民请直接勾选 No)", bold=True, size=12)
add_p("If you are a U.S. citizen, would you like to register to vote?   [  ] Yes   [  ] No")
add_p("If registered, would you like to update your voter information?   [  ] Yes   [  ] No")
add_p("Would you like to be an Election Judge?   [  ] Yes   [  ] No")

add_p("Emergency Contacts (紧急联系人 - 可选择性提供2位)", bold=True, size=12, space_after=8)
add_p("Contact 1 Name: _____________________  Phone: __________________________")
add_p("Address: ________________________________________________________")
add_p("Contact 2 Name: _____________________  Phone: __________________________")
add_p("Address: ________________________________________________________")

add_p("CERTIFICATION (申请人法定认证与签字)", bold=True, size=12, space_after=8)
add_p("I do solemnly swear, affirm, or certify that I am the person named herein and that the statements on this application are true and correct.")
add_p("1. I further certify my residence address is a (select one):")
add_p("    [  ] single family dwelling (独立屋)    [ X ] apartment (公寓)    [  ] motel    [  ] temporary shelter", bold=True)
add_p("2. I agree to immediately report to the Texas Department of Public Safety any changes in my medical condition which may affect my ability to safely operate a motor vehicle.")
add_p("3. I further understand that I am required by law to report any change of name or address to the Department of Public Safety within thirty days.")

p_sig = doc.add_paragraph()
p_sig.paragraph_format.space_before = Pt(24)
r_sig = p_sig.add_run("Signature of Applicant (申请人本人手写签名): _____________________________________\n\nDate (日期 MM/DD/YYYY): _____ / _____ / _________")
r_sig.font.name = 'Arial'
r_sig.font.size = Pt(11)
r_sig.bold = True

doc.save("DL-64_Form.docx")
print("🎉 成功！已在当前目录下为您生成 'DL-64_Form.docx' 文件！")